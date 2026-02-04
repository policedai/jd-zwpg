import streamlit as st
from openai import OpenAI
import base64
from io import BytesIO
from docx import Document
import re

# --- 1. 配置 API (适配云端安全版) ---
# 建议在 Streamlit Secrets 中配置 DASHSCOPE_API_KEY
client = OpenAI(
    api_key=st.secrets["DASHSCOPE_API_KEY"], 
    base_url="https://dashscope.aliyuncs.com/compatible-mode/v1",
    timeout=60.0
)

# --- 2. 工具函数 ---
def encode_image(image_file):
    return base64.b64encode(image_file.read()).decode('utf-8')

def create_docx(content, title):
    """格式化导出 Word 文档"""
    doc = Document()
    doc.add_heading(title, 0)
    for line in content.split('\n'):
        line = line.strip()
        if not line: continue
        if line.startswith('###'):
            doc.add_heading(line.replace('###', '').strip(), level=1)
        elif line.startswith('- **'):
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(line.replace('- ', '')).bold = True
        else:
            doc.add_paragraph(line)
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

# --- 3. 核心批改逻辑 ---
def grade_essay(image_base64, exam_type):
    if exam_type == "中考":
        score_detail = "满分 10 分（内容4/语言4/结构2）"
        word_requirement = "120词左右"
        grammar_rules = """
        【句式升级池】：
        - 复合句：灵活运用宾语从句、定语从句（that/who/which）或状语从句（unless, although, as soon as）。
        - 经典结构：使用 'It is + adj. + (for sb) to do'，'too... to...' 或 'so... that...'。
        - 比较结构：尝试 'The + 比较级, the + 比较级' 或 'one of the + 最高级'。
        【词汇与衔接】：
        - 词汇升级：用 extraordinary 代替 good, 用 essential 代替 important, 用 assist 代替 help。
        - 逻辑衔接：使用 Furthermore, However, Consequently, In a word 等地道连接词。
        """
        grading_focus = "侧重语言准确性、句式多样性及基础复合句的自然应用。"
    else:
        score_detail = "满分 20 分（内容8/语言8/结构4）"
        word_requirement = "120-140词"
        grammar_rules = """
        【高级语法池】：
        - 非谓语动词：使用现在分词(Doing)或过去分词(Done)作状语、定语或伴随状语。
        - 特殊句式：灵活嵌入【倒装句】（Only by.../Never have I...）、【强调句】（It is... that...）或【虚拟语气】（suggest that... should...）。
        - 复杂从句：尝试使用【主语从句】（What captures my attention is...）或【介词+which/whom】引导的定语从句。
        - With复合结构：'With + n. + doing/done/adj.' 增加描写生动性。
        【修辞与词汇】：
        - 词汇精准：使用 be convinced that (相信), fundamental (基础的), transform (改变) 等学术/高级词汇。
        - 逻辑修辞：使用排比(Parallelism)或感叹句增强感染力；衔接语使用 Admittedly, Paradoxically, To sum up。
        """
        grading_focus = "侧重立意高度、逻辑严密性以及对复杂句法结构和地道词汇的驾驭能力。"

    prompt = f"""
    你现在是北京市英语阅卷专家。请阅读图片中学生手写的英语作文。
    
    ### 任务要求：
    1. **识读提取**：提取题目要求与学生作文原文。
    2. **阅卷评分**：按北京{exam_type}{score_detail}标准评分。
    3. **精确字数统计**：
       - 请分别计算【学生原文】和【满分范文】的字数。
       - **规则**：仅计算单词数，严格剔除标点符号、空格及末尾落款。
    4. **逐句修改**：对习作原文进行深度解析。
       - 格式：【原句】 -> 【修改建议】 (从语法纠错、用词升级、或句式润色角度说明理由)。
    5. **深度点评**：{grading_focus}
    6. **满分范文**：写一篇{word_requirement}的满分范文。
       - **特别要求**：请根据题目语境，从以下【语法池】中**有机挑选 4-5 处高级表达**自然融入，严禁生搬硬套。
       {grammar_rules}
       - **范文亮点：将使用的这些高阶句式和词汇加粗。**

    输出格式：
    ---
    ### 1. 提取内容
    【题目内容】：...
    【学生原文】：...
    ### 2. {exam_type}阅卷结果
    - **【学生原文字数】**：XX 词 (不含标点)
    - **【最终得分】**：内容 X/语言 X/结构 X -> **总分 X**
    - **【逐句修改解析】**：
       - (原句1) -> (修改建议1) [理由]
       - (原句2) -> (修改建议2) [理由]
    - **【专家总评】**：...
    ### 3. 满分版修改作文
    - **【范文字数】**：XX 词 (不含标点)
    
    (范文正文，重点表达加粗)
    
    ### 4. 高阶句式解析
    (从范文中选出 2 个最值得学习的高级句式，讲解其语法结构与提分点)
    """

    try:
        response = client.chat.completions.create(
            model="qwen-vl-plus", 
            messages=[{"role": "user", "content": [
                {"type": "text", "text": prompt},
                {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{image_base64}"}}
            ]}]
        )
        return response.choices[0].message.content
    except Exception as e:
        return f"批改失败: {str(e)}"

# --- 4. Streamlit UI 界面 ---
st.set_page_config(page_title="北京中高考作文批改专家", layout="wide")

# 注入 CSS 优化间距
st.markdown("""
    <style>
        html, body, [class*="css"], .stMarkdown {
            font-size: 14px !important;
            line-height: 1.5 !important;
        }
        h1 { font-size: 1.8rem !important; color: #1E3A8A; }
        h2 { font-size: 1.4rem !important; border-bottom: 1px solid #ddd; padding-bottom: 5px; }
        h3 { font-size: 1.1rem !important; margin-top: 15px !important; color: #2563EB; }
        .block-container { padding-top: 2rem !important; }
        .stButton { margin-top: 10px; }
    </style>
    """, unsafe_allow_html=True)

st.title("📝 玛丽北京中高考英语作文批改")

if "report_text" not in st.session_state:
    st.session_state.report_text = None

with st.sidebar:
    st.header("🔐 访问授权")
    # 门禁系统：请在 Secrets 中设置 ACCESS_PASSWORD
    password = st.text_input("输入授权码", type="password")
    if password != st.secrets["ACCESS_PASSWORD"]:
        st.warning("请输入正确的授权码以解锁功能")
        st.stop()

    st.success("验证通过")
    st.write("---")
    st.header("⚙️ 阅卷设置")
    exam_choice = st.radio("选择评估标准", ["中考", "高考"])
    st.write("---")
    uploaded_file = st.file_uploader("上传作文照片", type=['jpg', 'jpeg', 'png'])
    
    if uploaded_file and st.button("🔄 重置报告"):
        st.session_state.report_text = None
        st.rerun()

col1, col2 = st.columns(2)

if uploaded_file:
    with col1:
        st.subheader("🖼️ 学生手稿")
        st.image(uploaded_file, use_container_width=True)
    
    with col2:
        st.subheader("📊 批改报告")
        
        if st.button("🚀 开始阅卷"):
            with st.spinner(f"正在调取北京{exam_choice}评分标准..."):
                img_b64 = encode_image(uploaded_file)
                st.session_state.report_text = grade_essay(img_b64, exam_choice)
        
        if st.session_state.report_text:
            tab_render, tab_raw = st.tabs(["✨ 可视化视图", "📄 Markdown 源码"])
            
            with tab_render:
                st.markdown(st.session_state.report_text)
            
            with tab_raw:
                st.code(st.session_state.report_text, language="markdown")
                
            st.write("---")
            doc_file = create_docx(st.session_state.report_text, f"北京{exam_choice}英语作文批改报告")
            st.download_button(
                label="📥 下载 Word 版报告",
                data=doc_file,
                file_name=f"北京{exam_choice}_批改报告.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
else:

    st.info("💡 请在左侧侧边栏上传作文图片（支持中考/高考标准）。")

